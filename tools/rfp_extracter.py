"""
RFP 파일의 타입에 따라 전처리 및 검색에 필요한 내용을 추출하는 모듈 모음
"""
import re
import io
import tempfile
import olefile
import zlib
import struct
from docx import Document

from dotenv import load_dotenv

load_dotenv("/workspace/Dhq_chatbot/pdf2DB_Agent/.env")

from google import genai
import pathlib
from google.genai.types import Part
import unicodedata

def clean_text(text):
    # 한자 및 보기 불편한 제어 문자 제거
    text = re.sub(r'[一-龥]', '', text)  # 한자 제거
    text = re.sub(r'[\x00-\x1F\x7F]', '', text)  # 제어문자 제거
    text = unicodedata.normalize("NFC", text)  # 정규화
    return text

gemini_client = genai.Client()

def extract_text_from_pdf(file_buffer, client=gemini_client):
    # file_buffer: BytesIO or bytes

    if hasattr(file_buffer, "read"):
        pdf_bytes = file_buffer.read()
    else:
        pdf_bytes = file_buffer

    prompt_for_ocr = """이는 PDF 문서 입니다. 구조를 유지하면서 모든 텍스트 콘텐츠를 추출하세요. 테이블, 열, 헤더 및 모든 구조화된 콘텐츠에 특별히 주의하세요. 단락 구분 및 형식을 유지하세요.

이 문서 페이지에서 모든 텍스트 콘텐츠를 추출하세요. 

`테이블의 경우:
1. 마크다운 테이블 형식을 사용하여 테이블 구조를 유지하세요.
2. 모든 열 머리글 및 행 레이블을 보존하세요.
3. 숫자 데이터를 정확하게 기입하세요.
4. 하나의 칸이 2개 이상의 칸과 매치되는 경우를 유의해서 추출하세요.`

`다중 열 레이아웃의 경우:
1. 왼쪽에서 오른쪽으로 열을 처리하세요.
2. 서로 다른 열의 콘텐츠를 명확하게 구분하세요.`

`차트 및 그래프의 경우:
1. 차트 유형을 설명하세요.
2. 가시적인 축 레이블, 범례 및 데이터 포인트를 추출하세요.
3. 제목 또는 캡션을 추출하세요.`"""

    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=[ Part.from_bytes(
            data=pdf_bytes,
            mime_type='application/pdf',
        ),
        prompt_for_ocr]
    )

    extracted_txt = response.text

    with tempfile.NamedTemporaryFile('w', suffix='.txt', dir='/tmp', delete=False, encoding='utf-8') as tmpf:
        tmpf.write(extracted_txt)
        tmp_txt_path = tmpf.name

    return tmp_txt_path

def extract_text_from_docx(file_buffer):

    if hasattr(file_buffer, "read"):
        file_buffer.seek(0)
        doc = Document(file_buffer)
    else:
        doc = Document(io.BytesIO(file_buffer))

    text = "\n".join([para.text for para in doc.paragraphs])

    with tempfile.NamedTemporaryFile('w', suffix='.txt', dir='/tmp', delete=False, encoding="utf-8") as tmpf:
        tmpf.write(text)
        tmp_txt_path = tmpf.name
    return tmp_txt_path

def extract_text_from_hwp(file_buffer):
    f = olefile.OleFileIO(file_buffer)
    dirs = f.listdir()

    if ["FileHeader"] not in dirs or ["\x05HwpSummaryInformation"] not in dirs:
        raise Exception("Not Valide HWP file")

    header = f.openstream("FileHeader")
    header_data = header.read()
    is_compressed = (header_data[36] & 1) == 1

    nums = []

    for d in dirs:
        if d[0] == "BodyText":
            nums.append(int(d[1][len("Section"):]))
    
    sections = ["BodyText/Section" + str(x) for x in sorted(nums)]

    text = ""

    for section in sections:
        bodytext = f.openstream(section)
        data = bodytext.read()
        if is_compressed:
            unpacked_data = zlib.decompress(data, -15)
        else:
            unpacked_data = data

        section_text = ""
        i = 0
        size = len(unpacked_data)

        while i < size:
            header = struct.unpack_from("<I", unpacked_data, i)[0]
            rec_type = header & 0x3ff
            rec_len = (header >> 20) & 0xfff

            if rec_type in [67]:
                rec_data = unpacked_data[i+4: i+4+rec_len]
                section_text += rec_data.decode('utf-16')
                section_text += "\n"

            i += 4 + rec_len

        text += section_text
        text += "\n"

        
    with tempfile.NamedTemporaryFile('w', suffix='.txt', dir='/tmp', delete=False, encoding="utf-8") as tmpf:
        tmpf.write(text)
        tmp_txt_path = tmpf.name
    return tmp_txt_path


def convert_rfp_for_RAG(tmp_txt_path, client=gemini_client):
    with open(tmp_txt_path, 'r', encoding='utf-8') as f:
        extracted_text = f.read()

#     prompt_for_convert = f"""아래의 제품 견적 요청서(RFP) 문서에서 필요한 제품의 스펙 요구 조건을 다음과 같은 규칙에 따라 추출하고 정리해주세요.

# **제품 스펙 요구 조건 추출 규칙**
# 1. 원본 파일에서 언급한 정보를 그대로 포함해야 합니다.
# 2. 추출한 정보는 아래의 작성 예시와 같이 key-value 형태로 작성해야 합니다.
# 3. 정보를 추출할 때, key 값은 제품의 정보 카테고리를 나타내며 value 값은 해당 카테고리에 대한 제품의 구체적인 정보를 나타냅니다.
# 4. key 값은 원본 파일의 제품 요구 조건을 설명하기에 적합해야 하며, 작성 예시와 같지 않아도 됩니다.

# **작성 예시**
# "product_name": "IBM Storage Scale",
# "storage_type": "소프트웨어 정의 스토리지",
# "supported_os": "IBM AIX, Linux (Red Hat, SUSE Linux Enterprise Server), Microsoft Windows Server 2012, Microsoft Windows 7, IBM z Systems",
# "protocols": "POSIX, GPFS, NFS v4.0, SMB v3.0",
# "big_data_support": "Hadoop MapReduce",
# "cloud_support": "OpenStack Cinder(블록), OpenStack Swift(오브젝트), S3(오브젝트)",
# "cloud_object_storage": "IBM Cloud Storage System (Cleversafe), Amazon S3, IBM Cloud Native Object, OpenStack Swift, Amazon S3 호환",
# "max_files_per_filesystem": "2의 64승(900경)개 파일",
# "max_filesystem_size": "8 엑사바이트(EB)",
# "max_data_capacity": "10억 페타바이트",
# "min_nodes": "1",
# "max_nodes": "16,384"
# ...


# [ RFP 문서 텍스트 ]
# {extracted_text}"""
    
#     response = client.models.generate_content(
#         model="gemini-2.5-flash",
#         contents=[prompt_for_convert]
#     )

#     rfp_search_term = response.text

    return extracted_text.strip()

